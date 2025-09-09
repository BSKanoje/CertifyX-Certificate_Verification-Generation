from datetime import date
from io import BytesIO
import os
import uuid

import qrcode
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.conf import settings


def generate_qr_image(data):
    qr = qrcode.make(data)
    return qr.convert('RGB')

def generate_certificate_pdf(template_obj, candidate):
    # Convert PDF template to image
    pages = convert_from_path(template_obj.file.path, dpi=300)
    template_img = pages[0]  # Take the first page

    # Save image to memory
    img_io = BytesIO()
    template_img.save(img_io, format='PNG')
    img_io.seek(0)
    background = ImageReader(img_io)

    # Create a new PDF
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=A4)

    # Draw background image
    c.drawImage(background, 0, 0, width=A4[0], height=A4[1])

    for field in template_obj.fields.all():
        field_key = field.field_name.strip().lower().replace(" ", "_")

        # Fetch value based on key name
        if field_key == 'name':
            value = candidate.name
        elif field_key == 'project_title':
            value = candidate.project_title  
        elif field_key == 'start_date':
            value = candidate.start_date.strftime('%B %d, %Y') if candidate.start_date else ''
        elif field_key == 'end_date':
            value = candidate.end_date.strftime('%B %d, %Y') if candidate.end_date else 'Ongoing'
        elif field_key == 'duration':
            if candidate.start_date:
                start = candidate.start_date.strftime('%B %d, %Y')
                end = candidate.end_date.strftime('%B %d, %Y') if candidate.end_date else 'Ongoing'
                value = f"{start} to {end}"
            else:
                value = ""
        elif field_key == 'issue_date':
            value = date.today().strftime('%B %d, %Y')
        else:
            value = getattr(candidate, field_key, '')

        # Set font and draw
        c.setFont(field.font, field.font_size)

        # Centering name only
        if field.field_name.lower() == 'name':
            text_width = c.stringWidth(value, field.font, field.font_size)
            center_x = (A4[0] - text_width) / 2
            c.drawString(center_x, field.position_y, value)
        else:
            c.drawString(field.position_x, field.position_y, value)

    # Generate UUID
    uid = str(uuid.uuid4())

    # Add Certificate ID to PDF (visible text)
    c.setFont("Helvetica", 10)
    c.drawString(50, 50, f"Certificate ID: {uid}")

    # Add QR code
    qr = qrcode.make(f"http://127.0.0.1:8000/generateCertificate/verify/?id={uid}")
    qr_io = BytesIO()
    qr.save(qr_io, format='PNG')
    qr_io.seek(0)
    c.drawImage(ImageReader(qr_io), 450, 25, width=100, height=100)

    # Add Company Logo (optional)
    if template_obj.company.logo:
        try:
            logo_path = os.path.join(settings.MEDIA_ROOT, template_obj.company.logo.name)
            logo = ImageReader(logo_path)
            c.drawImage(logo, x=50, y=A4[1] - 150, width=100, height=100, mask='auto')  # top-left corner
        except Exception as e:
            print("Logo error:", e)

    c.showPage()
    c.save()
    pdf_buffer.seek(0)

    return uid, pdf_buffer


def generate_certificate_docx(template_obj, candidate):
    doc = Document(template_obj.file.path)
    print("Candidate gender:", repr(candidate.gender))

    pronouns = get_pronouns(candidate.gender)


    # Replace placeholders
    replacements = {
        'NameOfTheStudent': candidate.name,
        'StudentUSN': getattr(candidate, 'usn', ''),
        'StudentDepartment': getattr(candidate, 'department', ''),
        'StudentCollege': getattr(candidate, 'college', ''),
        'StudentCompanyName': candidate.company.company_name if candidate.company else '',
        'StartDateOfInternship': candidate.start_date.strftime('%B %d, %Y') if candidate.start_date else '',
        'EndDateOfInternship': candidate.end_date.strftime('%B %d, %Y') if candidate.end_date else '',
        'InternshipProject': getattr(candidate, 'project_title', ''),
        **pronouns
    }

    replace_placeholders_preserving_format(doc, replacements)


    # Generate Certificate ID and QR Code
    certificate_id = str(uuid.uuid4())
    verification_url = f"http://127.0.0.1:8000/generateCertificate/verify/?id={certificate_id}"

    qr = qrcode.QRCode(box_size=2, border=1)
    qr.add_data(verification_url)
    qr.make(fit=True)
    qr_img = qr.make_image(fill='black', back_color='white')

    qr_io = BytesIO()
    qr_img.save(qr_io, format="PNG")
    qr_io.seek(0)

    # Add spacer paragraph
    doc.add_paragraph()

    # Add QR code (right-aligned)
    p_qr = doc.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_qr_run = p_qr.add_run()
    p_qr_run.add_picture(qr_io, width=Inches(0.8))

    # Add Certificate ID text (right-aligned)
    p_id = doc.add_paragraph()
    p_id.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_id = p_id.add_run(f"Certificate ID: {certificate_id}")
    run_id.font.size = Pt(8)

    # Save to DOCX in memory
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return certificate_id, docx_buffer


def get_pronouns(gender):
    if not gender:
        gender = "other"
    gender = gender.strip().lower()

    # Handle short codes
    if gender in ['m', 'male']:
        return {
            "he_she_they": "he",
            "him_her_them": "him",
            "his_her_their": "his",
            "He_She_They": "He",
            "Him_Her_Them": "Him",
            "His_Her_Their": "His",
        }
    elif gender in ['f', 'female']:
        return {
            "he_she_they": "she",
            "him_her_them": "her",
            "his_her_their": "her",
            "He_She_They": "She",
            "Him_Her_Them": "Her",
            "His_Her_Their": "Her",
        }
    else:
        return {
            "he_she_they": "they",
            "him_her_them": "them",
            "his_her_their": "their",
            "He_She_They": "They",
            "Him_Her_Them": "Them",
            "His_Her_Their": "Their",
        }


def replace_placeholders_preserving_format(doc, replacements):
    """
    Replace placeholders in a Word document while preserving all formatting.
    Works for paragraphs and tables.
    """
    def replace_in_paragraph(paragraph):
        for key, val in replacements.items():
            # Get all text in paragraph as a single string
            full_text = ''.join(run.text for run in paragraph.runs)
            if key not in full_text:
                continue

            # Replace placeholder in the full text
            new_text = full_text.replace(key, val)

            # Update runs to preserve formatting
            # Strategy: keep formatting of first run, clear all runs, assign new text to first run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ''

    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # Replace in all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
